"""
Generate MiniEnz manuscript DOCX — clean version with all 3 tables + 4 figures.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FIG_DIR = r"H:\eazyclaw\saved\MiniEnz_Methodology\manuscript\figures"
OUT_PATH = r"H:\eazyclaw\saved\MiniEnz_Methodology\reports\MiniEnz_Manuscript.docx"

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text):
    p = doc.add_paragraph(text)
    return p

def add_table_full(headers, data):
    """Add formatted table with header row."""
    tbl = doc.add_table(rows=1+len(data), cols=len(headers))
    tbl.style = 'Light List Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = tbl.cell(i+1, j)
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    return tbl

def add_figure(fname, caption):
    fpath = os.path.join(FIG_DIR, fname)
    if os.path.exists(fpath):
        img = doc.add_picture(fpath, width=Inches(5.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.size = Pt(9)
            r.italic = True

# ===== TITLE =====
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('MiniEnz: A Standardized Framework and Pilot Benchmark\nfor Deep Learning-Guided Enzyme Miniaturization')
r.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'

auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth.add_run('[Author names withheld for double-blind review]').font.size = Pt(10)

doc.add_paragraph()

# ===== ABSTRACT =====
add_heading('Abstract', 1)
add_para('Enzyme miniaturization—reducing protein size while preserving catalytic function—is a long-standing goal in protein engineering with applications spanning industrial biocatalysis, gene therapy cargo optimization, and biosensor development. Despite rapid advances in deep generative models for protein design, no standardized benchmark exists for evaluating computational enzyme miniaturization methods. We introduce MiniEnz, a standardized computational framework and pilot benchmark comprising 8 structurally diverse enzymes spanning 6 EC classes, 7 distinct fold types, and a 4.5x range in protein length. Evaluation through RFdiffusion (1,800 backbone scaffolds), ProteinMPNN (4,800 sequences), and ESM-2 embedding analysis revealed size reduction ranging from 6% to 50%, with 100% of designs maintaining sub-angstrom catalytic geometry. A key finding is that multi-motif scaffolding rescued pc_lipase from -20% to +50% SRE. We establish that catalytic RMSD and sequence design scores are independent evaluation dimensions (r ~ 0). ESM-2 analysis reveals three-tier ordering of embedding similarity: random > design-design > native-design. MiniEnz provides reusable evaluation metrics and establishes baseline benchmarks for computational enzyme miniaturization.')

kw = doc.add_paragraph()
kw.add_run('Keywords: ').bold = True
kw.add_run('enzyme miniaturization, protein design, RFdiffusion, ProteinMPNN, ESM-2, benchmark, deep learning').font.size = Pt(10)

# ===== 1. INTRODUCTION =====
add_heading('1. Introduction', 1)
add_heading('1.1 Why Miniaturize Enzymes?', 2)
add_para('Enzymes are nature\'s catalysts, but natural enzymes are often larger than functionally necessary. Reducing enzyme size offers higher expression yields, enhanced stability, faster diffusion, and AAV packaging compatibility. Industrial biocatalysis, gene therapy, and synthetic biology all demand compact, robust enzymes.')

add_heading('1.2 Deep Generative Models Transform Protein Design', 2)
add_para('RFdiffusion (Watson et al., 2023, Nature) and ProteinMPNN (Dauparas et al., 2022, Science) enable atomic-accuracy de novo design. The Baker lab achieved computational serine hydrolase design with kcat/Km up to 2.2 x 10^3 M-1 s-1 (Lauko et al., 2025, Science). However, no systematic study has evaluated general-purpose enzyme miniaturization across diverse families.')

add_heading('1.3 Contributions', 2)
for c in [
    'A curated pilot dataset of 8 enzymes (129–581 aa, 6 EC classes, 7 fold types)',
    'A reusable pipeline: RFdiffusion -> ProteinMPNN -> ESM-2 embedding analysis',
    'Size reduction 6–50%, 100% sub-angstrom catalytic RMSD, independent dual-dimension framework',
    'Multi-motif scaffolding discovery: pc_lipase improved from -20% to +50% SRE',
    'Systematic failure mode classification (Type A/B/C)',
    'Public release of all data (1,800 PDBs, 4,800 sequences, embeddings, code)',
]:
    doc.add_paragraph(c, style='List Bullet')

# ===== 2. METHODS =====
add_heading('2. Methods', 1)
add_heading('Table 1: MiniEnz Enzyme Catalog', 3)
add_table_full(
    ['Enzyme', 'PDB', 'EC', 'Length', 'Fold', 'Catalytic Residues', 'Cofactors'],
    [
        ['Lysozyme', '1LSE', '3.2.1.17', '129', 'alpha+beta', 'Glu35, Asp52', 'None'],
        ['Subtilisin BPN\'', '1SBT', '3.4.21.62', '275', 'alpha/beta sandwich', 'Asp32, His64, Ser221', 'None'],
        ['TEM-1 beta-Lactamase', '1BTL', '3.5.2.6', '263', 'alpha/beta DD-peptidase', 'Ser70, Lys73, Ser130, Glu166', 'None'],
        ['Triosephosphate Isomerase', '1TIM', '5.3.1.1', '247', 'TIM barrel', 'His95, Glu165', 'None'],
        ['Glucose Oxidase', '1GAL', '1.1.3.4', '581', 'FAD-binding', 'His516, His559', 'FAD'],
        ['P. cepacia Lipase', '3LIP', '3.1.1.3', '320', 'alpha/beta hydrolase', 'Ser87, Asp264, His286', 'None'],
        ['Carbonic Anhydrase II', '1CA2', '4.2.1.1', '256', 'All-beta', 'His94,96,119,Thr199,Glu106', 'Zn2+'],
        ['Tropinone Reductase-II', '2AE2', '1.1.1.184', '259', 'Rossmann SDR', 'Tyr155, Tyr159', 'NADP+'],
    ])
doc.add_paragraph()

add_heading('2.2 RFdiffusion Motif Scaffolding', 2)
add_para('RFdiffusion base model (T=50, linear beta-schedule) generated 200 backbone candidates per enzyme. Two strategies: single-motif (contiguous catalytic region fixed, 7 enzymes) and multi-motif (three independent short segments, pc_lipase). Total: 1,800 backbone PDB files.')

add_heading('2.3 ProteinMPNN Sequence Design', 2)
add_para('Top 50 backbones per enzyme (ranked by catalytic motif RMSD) through ProteinMPNN v_48_020 (Ca-only, T=0.2, 8 sequences/backbone). Fixed catalytic motif chains locked; scaffold chains designed. Total: 4,800 sequences.')

add_heading('2.4 ESM-2 Embedding Analysis', 2)
add_para('ESM-2 embeddings (esm2_t30_150M_UR50D, 640-dim) for all designed + 8 native wild-type + 200 random sequences. Three analyses: intra-design, cross-enzyme, and native-design similarity.')

# ===== 3. RESULTS =====
add_heading('3. Results', 1)

add_heading('3.1 Size Reduction Efficiency', 2)
add_figure('fig1_sre.png', 'Figure 1: Size Reduction Efficiency across 8 enzyme families. Bars show mean +/- SD from 200 designs.')

add_heading('Table 2: Benchmark Results', 3)
add_table_full(
    ['Enzyme', 'SRE (%)', 'MPNN Best', 'MPNN Mean +/- SD', 'Sim_intra', 'Sim_cross'],
    [
        ['P. cepacia Lipase', '50.2 +/- 4.4', '0.901', '1.129 +/- 0.080', '0.908', '0.943'],
        ['Glucose Oxidase', '38.4 +/- 5.9', '0.897', '1.045 +/- 0.053', '0.946', '0.955'],
        ['Carbonic Anhydrase II', '27.4 +/- 6.3', '0.897', '1.105 +/- 0.068', '0.925', '0.943'],
        ['Tropinone Reductase-II', '26.7 +/- 5.8', '0.889', '1.080 +/- 0.064', '0.937', '0.957'],
        ['Subtilisin BPN\'', '23.5 +/- 5.8', '0.927', '1.092 +/- 0.066', '0.936', '0.967'],
        ['Lysozyme', '20.4 +/- 5.4', '1.001', '1.206 +/- 0.100', '0.917', '0.934'],
        ['TEM-1 beta-Lactamase', '12.0 +/- 5.3', '0.956', '1.103 +/- 0.056', '0.940', '0.952'],
        ['Triosephosphate Isomerase', '6.0 +/- 6.1', '0.862', '1.031 +/- 0.061', '0.945', '0.958'],
        ['Random null baseline', '-', '-', '-', '0.967 +/- 0.019', '-'],
    ])
doc.add_paragraph()

add_heading('3.2 Active Site Geometry Preservation', 2)
add_para('All 1,800 designs maintained sub-angstrom catalytic Ca RMSD (<1.0 A). Mean RMSD ranged from 0.13 A (TEM-1, 4 residues) to 0.68 A (lysozyme, 2 residues). Catalytic RMSD showed no correlation with ProteinMPNN scores (r ~ 0).')

add_heading('Table 3: Active Site Geometry', 3)
add_table_full(
    ['Enzyme', 'Cat. Res.', 'Mean RMSD (A)', 'SD (A)', '<0.5 A (%)'],
    [
        ['TEM-1 beta-Lactamase', '4', '0.13', '0.04', '100.0'],
        ['Carbonic Anhydrase II', '5', '0.19', '0.05', '100.0'],
        ['Triosephosphate Isomerase', '2', '0.33', '0.05', '100.0'],
        ['P. cepacia Lipase', '3', '0.44', '0.15', '68.0'],
        ['Subtilisin BPN\'', '3', '0.50', '0.06', '48.0'],
        ['Glucose Oxidase', '2', '0.56', '0.07', '20.0'],
        ['Tropinone Reductase-II', '2', '0.62', '0.05', '2.0'],
        ['Lysozyme', '2', '0.68', '0.08', '2.0'],
    ])
doc.add_paragraph()

add_figure('fig5_cat_rmsd.png', 'Figure 2: Catalytic residue Ca RMSD distributions.')
add_figure('fig2_scores.png', 'Figure 3: ProteinMPNN score distributions.')

add_heading('3.3 ESM-2 Three-Way Comparison', 2)
add_para('Native-design similarity: 0.706 (lysozyme) to 0.843 (glucose oxidase), substantially below random baseline (0.967) and intra-design similarity (0.908–0.947). Three-tier ordering: random > design-design > native-design.')

add_figure('fig6_three_way_esm2.png', 'Figure 4: Three-way ESM-2 embedding comparison.')

add_heading('3.4 Failure Mode Analysis', 2)
add_para('Type A (Insufficient Compression): TIM barrel and DD-peptidase folds resist size reduction (6–12%). Type B (Contig Strategy Dependence): pc_lipase single-motif fails (SRE=-20%) but multi-motif succeeds (SRE=+50%). Type C (Cofactor Risk): Large-SRE enzymes may distort cofactor pockets, requiring RFdiffusion3 or docking validation.')

# ===== 4. DISCUSSION =====
add_heading('4. Discussion', 1)
add_para('Three factors predict miniaturization success: active site compactness (clustered > dispersed), fold architecture (all-beta > TIM barrel), and native size headroom. With n=8, these patterns are observational and warrant validation on a larger benchmark.')

add_heading('4.1 Limitations', 2)
add_para('(1) No AlphaFold2 validation due to ColabFold GPU incompatibility; ESM-2 provides sequence-level assessment only. (2) Single ProteinMPNN temperature (T=0.2). (3) No all-atom side-chain packing or cofactor docking. (4) No experimental validation—expression and activity assays needed for top candidates.')

# ===== 5. CONCLUSION =====
add_heading('5. Conclusion', 1)
add_para('MiniEnz demonstrates that deep generative models can achieve meaningful enzyme size reduction (6–50%) across diverse structural families, with 100% sub-angstrom catalytic geometry preservation. Multi-motif scaffolding dramatically improves dispersed active site miniaturization. Catalytic RMSD and sequence design scores are independent evaluation dimensions. All data and code are publicly available.')

# ===== DATA =====
add_heading('Data Availability', 1)
add_para('All RFdiffusion scaffolds (1,800 PDBs), ProteinMPNN sequences (4,800 FASTA), ESM-2 embeddings (9 NPZ files), analysis scripts, and evaluation code are available at [repository URL upon acceptance].')

add_heading('Code Availability', 1)
add_para('All tools used are open-source: RFdiffusion, ProteinMPNN, and ESM-2. Scripts for the full pipeline are provided in the repository.')

doc.save(OUT_PATH)
print("DOCX saved: %s" % OUT_PATH)
print("Tables: %d, Images: %d" % (len(doc.tables), sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)))
