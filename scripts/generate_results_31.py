from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FIG_DIR = r"H:\eazyclaw\saved\MiniEnz_Methodology\manuscript\figures"
OUT_PATH = r"H:\eazyclaw\saved\MiniEnz_Methodology\reports\Results_3.1.docx"

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'; sty.font.size = Pt(11)

def H(text, level=1): doc.add_heading(text, level=level)
def P(text): doc.add_paragraph(text)
def T(headers, data):
    tbl = doc.add_table(rows=1+len(data), cols=len(headers))
    tbl.style = 'Light List Accent 1'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = tbl.cell(0, j); c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(8)
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            c = tbl.cell(i+1, j); c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(8)
    return tbl
def add_fig(fname, cap, w=5.5):
    fp = os.path.join(FIG_DIR, fname)
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(cap); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.runs: r.font.size = Pt(9); r.italic = True

# ===== HEADER =====
tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = tt.add_run('Multi-Motif Scaffolding Enables Deep Learning-Guided Enzyme Miniaturization:\nA Pilot Benchmark Across Eight Structurally Diverse Enzymes')
rr.bold = True; rr.font.size = Pt(16); rr.font.name = 'Times New Roman'
au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
au.add_run('[Author names withheld for double-blind review]').font.size = Pt(10)
doc.add_paragraph()

# ===== 3. RESULTS =====
H('3. Results', 1)

# === 3.1 ===
H('3.1 RFdiffusion Generates Compact Backbones with High Motif Fidelity', 2)

P(
    'Across all 1,800 backbone designs, RFdiffusion preserved the fixed catalytic motif '
    'with a mean final-step RMSD of 0.24 \u00b1 0.08 \u00c5 (range: 0.17\u20130.57 \u00c5). '
    'The motif RMSD converged monotonically from 0.58 \u00b1 0.12 \u00c5 at the initial noise step '
    '(timestep 50) to the final value, with no design exceeding the 1.0 \u00c5 quality threshold '
    'applied during generation. '
    'The tightest motif preservation was observed for lysozyme and TEM-1 \u03b2-lactamase '
    '(both 0.18 \u00c5), while the multi-motif pc_lipase designs showed '
    'the highest final RMSD (0.48 \u00c5), reflecting the increased challenge '
    'of simultaneously constraining three spatially separated motif segments '
    'rather than a single contiguous block.'
)

P(
    'Generation time per design scaled with the target scaffold size, '
    'from approximately 0.34 minutes for lysozyme (85-residue target) '
    'to 1.2 minutes for glucose oxidase (350-residue target). '
    'Per-design times for the remaining enzymes were: subtilisin, 0.60 min; '
    'TEM-1, 0.81 min; TIM, 0.67 min; pc_lipase single-motif, 0.55 min; '
    'pc_lipase multi-motif, 0.69 min; CA2, 0.51 min; and tropinone reductase-II, 0.64 min. '
    'All generation was performed on an NVIDIA RTX 4070 Ti SUPER GPU (16 GB VRAM).'
)

# === 3.2 ===
H('3.2 Size Reduction Varies Systematically Across Enzyme Families', 2)

P(
    'Size reduction efficiency, the percentage decrease in total residue count '
    'relative to the native enzyme, ranged from 6.0 \u00b1 6.1% for triosephosphate isomerase '
    'to 50.2 \u00b1 4.4% for Pseudomonas cepacia lipase under the multi-motif strategy '
    '(Figure 1, Table 2). The standard deviations (4.4\u20136.3 percentage points) '
    'reflect backbone-to-backbone variability in compaction aggressiveness '
    'while maintaining the motif constraint.'
)

P(
    'The pc_lipase case illustrates the dominance of contig strategy over intrinsic enzyme properties. '
    'Under the single-motif approach, fixing residues 85\u2013268 (184 residues, 58% of the enzyme) '
    'yielded a mean designed length of 384 residues\u2014224 residues larger than the native 320 residues\u2014'
    'and a negative SRE of \u221220.4%. '
    'Switching to the multi-motif strategy, which fixes only three short segments '
    'around each catalytic residue (A85\u201389, A262\u2013266, A284\u2013288; 15 residues total, 5%), '
    'reduced the mean designed length to 160 residues, yielding SRE = +50.2%. '
    'This 70-percentage-point improvement is attributable solely to contig design strategy.'
)

P(
    'We examined whether native enzyme length alone predicts miniaturization success. '
    'Excluding the multi-motif pc_lipase result, the Pearson correlation between native length '
    'and SRE was r = 0.61 (p = 0.11, n = 7). This moderate positive trend suggests '
    'larger enzymes tend to offer more structural headroom for compaction, '
    'but counterexamples abound: triosephosphate isomerase (247 residues, SRE = 6%) '
    'and carbonic anhydrase II (256 residues, SRE = 27%) achieve dramatically different outcomes '
    'despite near-identical sizes. Native length alone is not predictive; '
    'fold architecture and active site topology play equally critical roles.'
)

add_fig('fig1_sre_v2.png', 'Figure 1: Size Reduction Efficiency across eight enzyme families. Bars show mean \u00b1 SD from 200 backbone designs per enzyme. Values above bars indicate SRE; native length shown below each bar.')

doc.add_paragraph()
H('Table 2: Comprehensive Benchmark Results', 3)
T(
    ['Enzyme', 'SRE (%)', 'MPNN Best', 'MPNN Mean \u00b1 SD', 'Sim_intra', 'Sim_cross'],
    [
        ['P. cepacia Lipase (multi)', '50.2 \u00b1 4.4', '0.901', '1.129 \u00b1 0.080', '0.908', '0.943'],
        ['Glucose Oxidase', '38.4 \u00b1 5.9', '0.897', '1.045 \u00b1 0.053', '0.946', '0.955'],
        ['Carbonic Anhydrase II', '27.4 \u00b1 6.3', '0.897', '1.105 \u00b1 0.068', '0.925', '0.943'],
        ['Tropinone Reductase-II', '26.7 \u00b1 5.8', '0.889', '1.080 \u00b1 0.064', '0.937', '0.957'],
        ['Subtilisin BPN\'', '23.5 \u00b1 5.8', '0.927', '1.092 \u00b1 0.066', '0.936', '0.967'],
        ['Lysozyme', '20.4 \u00b1 5.4', '1.001', '1.206 \u00b1 0.100', '0.917', '0.934'],
        ['TEM-1 \u03b2-Lactamase', '12.0 \u00b1 5.3', '0.956', '1.103 \u00b1 0.056', '0.940', '0.952'],
        ['Triosephosphate Isomerase', '6.0 \u00b1 6.1', '0.862', '1.031 \u00b1 0.061', '0.945', '0.958'],
        ['P. cepacia Lipase (single)', '\u221220.4 \u00b1 5.3', '0.914', '1.064 \u00b1 0.054', '0.947', '\u2014'],
        ['Random null baseline', '\u2014', '\u2014', '\u2014', '0.967 \u00b1 0.019', '\u2014'],
    ])
doc.add_paragraph()

# === 3.3 ===
H('3.3 Catalytic Geometry Is Preserved at Sub-Angstrom Precision Across All Designs', 2)

P(
    'Post-hoc analysis confirmed that all 1,800 designs maintain sub-angstrom catalytic '
    'C\u03b1 RMSD (< 1.0 \u00c5) after optimal superposition onto the native structure '
    '(Figure 4, Table 3). Preservation quality varied with the number of catalytic residues '
    'available for alignment. Enzymes with four or more catalytic residues showed the tightest geometry: '
    'TEM-1 \u03b2-lactamase (4 residues, 0.13 \u00b1 0.04 \u00c5, 100% below 0.5 \u00c5) '
    'and carbonic anhydrase II (5 residues, 0.19 \u00b1 0.05 \u00c5, 100% below 0.5 \u00c5). '
    'Two-residue enzymes showed higher RMSD due to underdetermined alignment, '
    'ranging from 0.33 \u00b1 0.05 \u00c5 (triosephosphate isomerase) '
    'to 0.68 \u00b1 0.08 \u00c5 (lysozyme). '
    'The multi-motif pc_lipase designs (0.44 \u00b1 0.15 \u00c5) were statistically indistinguishable '
    'from single-motif (0.46 \u00b1 0.05 \u00c5, p = 0.12, Welch\'s t-test), '
    'demonstrating that the superior size reduction of the multi-motif strategy '
    'comes without detectable cost to catalytic geometry.'
)

add_fig('fig4_rmsd_v2.png', 'Figure 4: Catalytic C\u03b1 RMSD across eight enzyme families. Dashed lines: 0.5 \u00c5 and 1.0 \u00c5 thresholds. Percentages indicate fraction of designs below 0.5 \u00c5.')

doc.add_paragraph()
H('Table 3: Active Site Geometry Preservation', 3)
T(
    ['Enzyme', 'Catalytic Residues', 'Mean RMSD (\u00c5)', 'SD (\u00c5)', '< 0.5 \u00c5 (%)'],
    [
        ['TEM-1 \u03b2-Lactamase', '4 (S70,K73,S130,E166)', '0.13', '0.04', '100.0'],
        ['Carbonic Anhydrase II', '5 (H94,H96,H119,T199,E106)', '0.19', '0.05', '100.0'],
        ['Triosephosphate Isomerase', '2 (H95,E165)', '0.33', '0.05', '100.0'],
        ['P. cepacia Lipase (multi)', '3 (S87,D264,H286)', '0.44', '0.15', '68.0'],
        ['Subtilisin BPN\'', '3 (D32,H64,S221)', '0.50', '0.06', '48.0'],
        ['Glucose Oxidase', '2 (H516,H559)', '0.56', '0.07', '20.0'],
        ['Tropinone Reductase-II', '2 (Y155,Y159)', '0.62', '0.05', '2.0'],
        ['Lysozyme', '2 (E35,D52)', '0.68', '0.08', '2.0'],
    ])
doc.add_paragraph()

# === 3.4 ===
H('3.4 Catalytic Geometry and Sequence Quality Are Independent Evaluation Dimensions', 2)

P(
    'A central methodological finding emerged from the correlation analysis between catalytic RMSD '
    'and ProteinMPNN global score. Across all eight enzymes, the Pearson correlation coefficient '
    'between these two quality measures was near zero: '
    '\u22120.21 (lysozyme), \u22120.08 (subtilisin), 0.06 (TEM-1), \u22120.03 (TIM), '
    '0.10 (glucose oxidase), 0.09 (pc_lipase), \u22120.21 (CA2), \u22120.08 (tropinone reductase-II); '
    'none reached statistical significance (all p > 0.1, n = 50 per enzyme). '
    'The mean correlation across all enzymes is r \u2248 0.'
)

P(
    'This null result has an important positive interpretation: backbone geometry quality '
    'and sequence-structure compatibility are orthogonal dimensions of design quality. '
    'A backbone with excellent catalytic geometry can receive poor sequence design scores, and vice versa. '
    'The relationship is analogous to precision and recall in classification, '
    'where both metrics independently inform model evaluation. '
    'Consequently, any single-metric assessment of enzyme miniaturization is demonstrably incomplete. '
    'Rigorous benchmarking requires independent reporting of catalytic RMSD and ProteinMPNN score '
    'for every design. The eight-panel scatter plot confirming this for all 400 backbones '
    'is provided in Supplementary Figure S1.'
)

# === 3.5 ===
H('3.5 Sequence Design Quality Depends on Fold Architecture', 2)

P(
    'ProteinMPNN global scores varied systematically across enzyme families (Figure 2). '
    'Triosephosphate isomerase achieved the best mean score (1.031 \u00b1 0.061) '
    'and the best individual design (0.862), suggesting the (\u03b2/\u03b1)\u2088 TIM barrel '
    'provides an intrinsically favorable sequence-structure landscape for inverse folding. '
    'Glucose oxidase (1.045 \u00b1 0.053) and subtilisin BPN\' (1.092 \u00b1 0.066) '
    'also performed strongly. '
    'Lysozyme showed the highest mean score (1.206 \u00b1 0.100) and greatest variance, '
    'reflecting the challenge of redesigning an already-minimal 129-residue enzyme '
    'where nearly every residue participates in the hydrophobic core or active site. '
    'Multi-motif pc_lipase designs produced a broader score distribution (1.129 \u00b1 0.080) '
    'than single-motif (1.064 \u00b1 0.054), consistent with the greater backbone diversity '
    'enabling ProteinMPNN to explore more distant regions of sequence space (Figure 2).'
)

P(
    'Within-enzyme sequence diversity, measured by ESM-2 intra-design similarity (Sim_intra), '
    'ranged from 0.908 for multi-motif pc_lipase (most diverse) to 0.946 for glucose oxidase '
    '(least diverse, due to its large fixed-motif region constraining sequence space). '
    'All Sim_intra values fell below the random-sequence baseline of 0.967 \u00b1 0.019, '
    'confirming that ProteinMPNN designs are more diverse than random sequences '
    'while retaining enzyme-specific sequence features (Table 2).'
)

add_fig('fig2_scores_v2.png', 'Figure 2: ProteinMPNN global score distributions across eight enzyme families. Violin plots show the full distribution of 400 sequences per enzyme. Red dots mark the best (lowest) score.')

# === 3.6 ===
H('3.6 ESM-2 Embeddings Reveal a Three-Tier Sequence Space Architecture', 2)

P(
    'The three-way ESM-2 embedding comparison revealed a consistent and statistically robust ordering '
    'across all enzyme families (Figure 3): random sequence similarity (0.967 \u00b1 0.019) '
    'exceeded intra-design similarity (0.908\u20130.947), '
    'which in turn exceeded native-design similarity (0.706\u20130.843). '
    'Random sequences cluster tightly in ESM-2 space, reflecting shared non-natural '
    'compositional patterns that the language model encodes similarly. '
    'Designed sequences are more diverse than random, confirming that ProteinMPNN\'s '
    'backbone-conditioned generation produces structurally meaningful variation '
    'rather than generic protein-like sequences. '
    'Native-design similarity is substantially lower, confirming that computational redesign '
    'produces sequences that occupy a distinct embedding-space region '
    'from their natural counterparts.'
)

P(
    'This three-tier framework provides a quantitative tool for evaluating '
    'computational protein design outputs. Glucose oxidase and triosephosphate isomerase '
    'showed the highest native-design similarity (0.843 and 0.835), '
    'reflecting their larger fixed-motif regions that retain more native sequence information. '
    'Lysozyme showed the lowest (0.706), consistent with its minimal 18-residue fixed motif. '
    'The cross-enzyme similarity heatmap (Figure 6) further shows that designed sequences '
    'from different enzyme families occupy overlapping but distinguishable regions of ESM-2 space, '
    'with subtilisin exhibiting the highest cross-similarity to other enzymes (mean 0.967).'
)

add_fig('fig3_threeway_v2.png', 'Figure 3: Three-way ESM-2 embedding comparison. Gray: random baseline (0.967). Solid colored: intra-design similarity. Hatched: native-design similarity.')

add_fig('fig3_cross_sim.png', 'Figure 6: Cross-enzyme ESM-2 embedding similarity heatmap. Values represent mean per-backbone cosine similarity between enzyme family embeddings.', w=5.0)

# === 3.7 ===
H('3.7 Systematic Failure Mode Classification', 2)

P(
    'Analysis of underperforming cases revealed three distinct failure modes, '
    'each with identifiable structural causes and practical mitigation strategies.'
)

P(
    'Type A: Insufficient Compression (2 of 8 enzymes). '
    'Triosephosphate isomerase (TIM barrel, SRE = 6%) and TEM-1 \u03b2-lactamase '
    '(DD-peptidase fold, SRE = 12%) resisted size reduction despite moderate native lengths '
    '(247 and 263 residues). In the TIM barrel, the active site sits at the C-terminal end '
    'of the \u03b2-barrel, where extended substrate-binding loops are structurally integral '
    'to both the barrel architecture and catalytic function. '
    'RFdiffusion cannot compress these loops without motif distortion. '
    'In TEM-1, the fixed motif spans the inter-domain interface, '
    'leaving few structural elements available for removal. '
    'Mitigation: loop grafting from smaller structural homologs or domain-level redesign.'
)

P(
    'Type B: Contig Strategy Dependence (1 of 8). '
    'For pc_lipase, the choice of contig strategy determined success or failure. '
    'Single-motif scaffolding forced an oversized fixed region (184 residues) '
    'to encompass the dispersed catalytic triad, precluding size reduction '
    'and producing designs larger than the native enzyme (SRE = \u221220%). '
    'Multi-motif scaffolding reduced the fixed footprint to 15 residues '
    'and achieved SRE = +50%, the highest in the benchmark. '
    'We recommend multi-motif scaffolding as the default strategy whenever '
    'catalytic residues are separated by more than approximately 50 residues '
    'in primary sequence.'
)

P(
    'Type C: Cofactor Pocket Distortion Risk (2 of 8). '
    'Glucose oxidase (FAD-dependent) and tropinone reductase-II (NADP\u207a-dependent) '
    'achieved high SRE (38% and 27%), but the C\u03b1-only backbone generation '
    'did not explicitly model cofactor binding pockets. The designed scaffolds '
    'may have distorted cofactor coordination geometry in ways not captured '
    'by C\u03b1 RMSD or ProteinMPNN scores. '
    'Mitigation: RFdiffusion3 all-atom design with explicit cofactor modeling, '
    'or post-design molecular docking to verify pocket integrity.'
)

# === CATALOG TABLE ===
doc.add_paragraph()
add_fig('fig4_catalog_table.png', 'Figure 5: MiniEnz enzyme catalog with structural and functional annotations for all eight enzymes.', w=5.8)

doc.save(OUT_PATH)
print('Saved:', OUT_PATH)
