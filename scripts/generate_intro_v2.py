from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_PATH = r"H:\eazyclaw\saved\MiniEnz_Methodology\reports\Introduction_V2.docx"

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'; sty.font.size = Pt(11)

def H(text, level=1):
    doc.add_heading(text, level=level)

def P(text):
    doc.add_paragraph(text)

def add_table(headers, data):
    tbl = doc.add_table(rows=1+len(data), cols=len(headers))
    tbl.style = 'Light List Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = tbl.cell(0, j); c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(8)
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            c = tbl.cell(i+1, j); c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(8)
    return tbl

# ===== TITLE =====
tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = tt.add_run('Multi-Motif Scaffolding Enables Deep Learning-Guided Enzyme Miniaturization:\nA Pilot Benchmark Across Eight Structurally Diverse Enzymes')
rr.bold = True; rr.font.size = Pt(16); rr.font.name = 'Times New Roman'

au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
au.add_run('[Author names withheld for double-blind review]').font.size = Pt(10)
doc.add_paragraph()

# ===== INTRODUCTION (3 paragraphs, no subheadings) =====
H('1. Introduction', 1)

# Paragraph 1: The problem — why miniaturize, traditional limitations, the computational opportunity
P(
    'Enzymes are nature\'s catalysts, driving chemical transformations with unparalleled specificity '
    'under mild conditions. Yet natural enzymes are rarely minimal. Many carry structural elaborations '
    'acquired through evolution\u2014domain insertions, extended loops, oligomerization interfaces\u2014'
    'that are not required for catalysis. Removing these non-essential elements '
    'offers clear practical advantages: smaller proteins typically express at higher yields in heterologous systems, '
    'resist thermal denaturation more robustly, diffuse faster through biosensor matrices, '
    'and, critically, fit within the packaging limits of adeno-associated virus vectors '
    '(approximately 4.7 kb), a bottleneck for gene therapy applications [1\u20136]. '
    'Experimentally, enzyme miniaturization has been pursued through rational domain truncation, '
    'random deletion screening, and directed evolution starting from thermostable orthologs [7\u20139]. '
    'These efforts have produced notable successes\u2014a 182-residue miniaturized \u03b2-lactamase '
    'retaining 40% catalytic activity, a 100-residue chorismate mutase maintaining full efficiency '
    '[10,11]\u2014but the underlying approach remains case-by-case. '
    'Most truncations abolish activity, with success rates typically below 5% [12]. '
    'The field has lacked a systematic, transferable method for predicting which structural elements '
    'can be removed from which enzymes without destroying catalytic function.'
)

# Paragraph 2: The computational transformation — RFdiffusion, ProteinMPNN, Baker lab, the knowledge gap
P(
    'The past three years have transformed this landscape. '
    'RFdiffusion, a denoising diffusion probabilistic model trained on protein backbone structures, '
    'can generate novel protein backbones conditioned on diverse structural constraints including '
    'active site motifs, binding interfaces, and symmetry requirements [13]. '
    'ProteinMPNN, an inverse folding model based on message-passing neural networks, '
    'designs amino acid sequences that fold into specified backbone structures with high accuracy [14]. '
    'Together, these tools have enabled de novo design of protein binders [15], symmetric oligomers [16], '
    'and catalytic sites embedded in novel scaffolds [17]. '
    'The Baker laboratory recently demonstrated that RFdiffusion, combined with ensemble-based active site '
    'preorganization assessment, can design serine hydrolases with catalytic efficiencies '
    '(k\u1d2c\u1d2c/K\u2098) reaching 2.2 \u00d7 10\u00b3 M\u207b\u00b9s\u207b\u00b9 '
    'starting from minimal active site descriptions [17]. '
    'That work established that deep generative models can embed catalytic machinery into '
    'protein scaffolds of varying sizes. However, it focused on building catalytic function '
    'into new scaffolds\u2014an approach we call "building up"\u2014rather than systematically '
    'reducing the size of existing enzymes, or "trimming down." '
    'Whether RFdiffusion-based pipelines can serve as general-purpose tools for enzyme miniaturization '
    'across diverse protein families has remained an open question.'
)

# Paragraph 3: The gap, what we do, contributions (no subsection labels)
P(
    'This question has persisted because the field lacks any standardized benchmark for evaluating '
    'computational enzyme miniaturization. Existing resources address adjacent but distinct problems: '
    'ProteinGym benchmarks mutation fitness prediction [18], FLIP evaluates protein engineering tasks [19], '
    'and a recent comprehensive review catalogued miniaturization strategies without providing '
    'an evaluation dataset [12]. None of these captures the multi-objective nature of miniaturization\u2014'
    'the simultaneous requirements for size reduction, catalytic geometry preservation, '
    'sequence quality, and structural compactness. Systematic documentation of failure modes\u2014'
    'what does not work and why\u2014is notably absent from the computational protein design literature. '
    'Here we introduce MiniEnz, a computational framework and pilot benchmark that addresses this gap. '
    'We selected 8 structurally and functionally diverse enzymes spanning six EC classes, '
    'seven fold types, and a 4.5-fold range in protein length (129 to 581 residues; Table 1). '
    'Each enzyme was processed through a three-stage pipeline: RFdiffusion motif scaffolding '
    'with two complementary strategies (single-motif and multi-motif), ProteinMPNN sequence design '
    'with chain-specific constraints, and ESM-2 protein language model embedding analysis '
    'within a novel three-way comparison framework. Our principal findings are that '
    'catalytic geometry and sequence design quality are statistically independent evaluation dimensions '
    '(r \u2248 0 across all enzymes), that multi-motif scaffolding\u2014fixing only catalytic residues '
    'as short independent segments rather than a single contiguous block\u2014can rescue miniaturization '
    'for enzymes with dispersed active sites (improving size reduction by 70 percentage points for '
    'a bacterial lipase), and that designed sequences occupy a distinct embedding-space region '
    'between random and native sequences. '
    'We publicly release all computational data: 1,800 backbone scaffolds, 4,800 designed sequences, '
    'ESM-2 embeddings, and analysis code.'
)

# ===== TABLE 1: ENZYME CATALOG =====
doc.add_paragraph()
doc.add_heading('Table 1: MiniEnz Enzyme Catalog', 3)
add_table(
    ['Enzyme', 'PDB', 'EC', 'Length', 'Fold', 'Catalytic Residues', 'Cofactors'],
    [
        ['Lysozyme', '1LSE', '3.2.1.17', '129', '\u03b1+\u03b2', 'Glu35, Asp52', 'None'],
        ['Subtilisin BPN\'', '1SBT', '3.4.21.62', '275', '\u03b1/\u03b2 sandwich', 'Asp32, His64, Ser221', 'None'],
        ['TEM-1 \u03b2-Lactamase', '1BTL', '3.5.2.6', '263', '\u03b1/\u03b2 DD-peptidase', 'Ser70, Lys73, Ser130, Glu166', 'None'],
        ['Triosephosphate Isomerase', '1TIM', '5.3.1.1', '247', 'TIM barrel', 'His95, Glu165', 'None'],
        ['Glucose Oxidase', '1GAL', '1.1.3.4', '581', 'FAD-binding', 'His516, His559', 'FAD'],
        ['P. cepacia Lipase', '3LIP', '3.1.1.3', '320', '\u03b1/\u03b2 hydrolase', 'Ser87, Asp264, His286', 'None'],
        ['Carbonic Anhydrase II', '1CA2', '4.2.1.1', '256', 'All-\u03b2', 'His94,96,119,Thr199,Glu106', 'Zn\u00b2\u207a'],
        ['Tropinone Reductase-II', '2AE2', '1.1.1.184', '259', 'Rossmann SDR', 'Tyr155, Tyr159', 'NADP\u207a'],
    ])

doc.save(OUT_PATH)
print('Saved:', OUT_PATH)
