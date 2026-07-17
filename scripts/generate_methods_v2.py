from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_PATH = r"H:\eazyclaw\saved\MiniEnz_Methodology\reports\Methods_V2.docx"

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'; sty.font.size = Pt(11)

def H(text, level=1):
    doc.add_heading(text, level=level)

def P(text):
    doc.add_paragraph(text)

def add_table(headers, data):
    tbl = doc.add_table(rows=1+len(data), cols=len(headers))
    tbl.style = 'Light List Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = tbl.cell(0, j); c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(8)
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            c = tbl.cell(i+1, j); c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(8)
    return tbl

# ===== TITLE =====
tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = tt.add_run('Multi-Motif Scaffolding Enables Deep Learning-Guided Enzyme Miniaturization:\nA Pilot Benchmark Across Eight Structurally Diverse Enzymes')
rr.bold = True; rr.font.size = Pt(16); rr.font.name = 'Times New Roman'

au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
au.add_run('[Author names withheld for double-blind review]').font.size = Pt(10)
doc.add_paragraph()

# ===== 2. METHODS =====
H('2. Methods', 1)

# === 2.1 Enzyme Selection and Dataset Curation ===
H('2.1 Enzyme Selection and Dataset Curation', 2)

P(
    'We selected 8 enzymes to maximize structural and functional diversity along four axes. '
    'First, we required coverage of Enzyme Commission (EC) classes with industrial or biomedical relevance, '
    'spanning oxidoreductases (EC 1), hydrolases (EC 3), lyases (EC 4), and isomerases (EC 5). '
    'Second, we selected enzymes representing distinct CATH fold architectures: all-\u03b1, all-\u03b2, '
    '\u03b1/\u03b2, \u03b1+\u03b2, (\u03b2/\u03b1)\u2088 TIM barrel, Rossmann fold, '
    'and FAD-binding domain. Third, we sought a gradient of active site complexity, '
    'from simple two-residue acid-base pairs (lysozyme: Glu35-Asp52) through catalytic triads '
    '(subtilisin: Asp32-His64-Ser221) to a five-residue Zn\u00b2\u207a-dependent metalloenzyme '
    '(carbonic anhydrase II) and cofactor-dependent enzymes (glucose oxidase with FAD, '
    'tropinone reductase-II with NADP\u207a). Fourth, all selected enzymes possess '
    'high-resolution crystal structures in the Protein Data Bank (PDB) with resolutions '
    'of 2.5 \u00c5 or better, ensuring reliable atomic coordinates for motif definition. '
    'We excluded oligomeric enzymes whose catalytic activity depends on quaternary structure, '
    'and enzymes whose active sites span multiple domains in a manner that '
    'precludes domain-level miniaturization through our motif-scaffolding approach.'
)
P(
    'The final catalog (Table 1) includes hen egg lysozyme (PDB: 1LSE, 129 residues, EC 3.2.1.17), '
    'subtilisin BPN\' from Bacillus amyloliquefaciens (1SBT, 275 residues, EC 3.4.21.62), '
    'TEM-1 \u03b2-lactamase from Escherichia coli (1BTL, 263 residues, EC 3.5.2.6), '
    'triosephosphate isomerase from Gallus gallus (1TIM, 247 residues, EC 5.3.1.1), '
    'glucose oxidase from Aspergillus niger (1GAL, 581 residues, EC 1.1.3.4), '
    'Pseudomonas cepacia lipase (3LIP, 320 residues, EC 3.1.1.3), '
    'human carbonic anhydrase II (1CA2, 256 residues, EC 4.2.1.1), '
    'and tropinone reductase-II from Datura stramonium (2AE2, 259 residues, EC 1.1.1.184). '
    'For each enzyme, catalytic residues were identified from the primary literature '
    'and verified in the PDB structure. For triosephosphate isomerase, which crystallizes as a homodimer, '
    'only chain A was used. All PDB files were downloaded from the RCSB Protein Data Bank '
    'and used without further modification.'
)

doc.add_paragraph()
doc.add_heading('Table 1: MiniEnz Enzyme Catalog', 3)
add_table(
    ['Enzyme', 'PDB', 'EC', 'Length (aa)', 'Fold Type', 'Catalytic Residues', 'Cofactors'],
    [
        ['Lysozyme', '1LSE', '3.2.1.17', '129', '\u03b1+\u03b2 (lysozyme-like)', 'Glu35, Asp52', 'None'],
        ['Subtilisin BPN\'', '1SBT', '3.4.21.62', '275', '\u03b1/\u03b2 sandwich (subtilisin-like)', 'Asp32, His64, Ser221', 'None'],
        ['TEM-1 \u03b2-Lactamase', '1BTL', '3.5.2.6', '263', '\u03b1/\u03b2 DD-peptidase', 'Ser70, Lys73, Ser130, Glu166', 'None'],
        ['Triosephosphate Isomerase', '1TIM', '5.3.1.1', '247', '(\u03b2/\u03b1)\u2088 TIM barrel', 'His95, Glu165', 'None'],
        ['Glucose Oxidase', '1GAL', '1.1.3.4', '581', 'FAD-binding + substrate domain', 'His516, His559', 'FAD'],
        ['P. cepacia Lipase', '3LIP', '3.1.1.3', '320', '\u03b1/\u03b2 hydrolase', 'Ser87, Asp264, His286', 'None'],
        ['Carbonic Anhydrase II', '1CA2', '4.2.1.1', '256', 'All-\u03b2 (10-strand twisted sheet)', 'His94, His96, His119, Thr199, Glu106', 'Zn\u00b2\u207a'],
        ['Tropinone Reductase-II', '2AE2', '1.1.1.184', '259', 'Rossmann fold (SDR family)', 'Tyr155, Tyr159', 'NADP\u207a'],
    ])

doc.add_paragraph()

# === 2.2 RFdiffusion Motif Scaffolding ===
H('2.2 RFdiffusion Motif Scaffolding', 2)

P(
    'For each enzyme, we defined a structural motif comprising all catalytic residues identified in the literature '
    'plus 2\u20133 flanking residues on each side. These flanking residues provide local structural context '
    'that helps RFdiffusion maintain proper backbone geometry around the catalytic site. '
    'The motif residues were specified as fixed constraints in RFdiffusion\u2019s motif scaffolding mode, '
    'while the surrounding protein scaffold was free to diffuse into a more compact structure. '
    'We employed two complementary contig strategies depending on the spatial distribution '
    'of catalytic residues in the primary sequence.'
)

P(
    'For seven enzymes with clustered catalytic residues, we used a single-motif strategy. '
    'The entire contiguous region spanning the catalytic residues is specified as a fixed receptor block '
    'in the RFdiffusion contig format: [chain_id start\u2013end/0 len_min\u2013len_max]. '
    'The first segment (chain_id start\u2013end/0) designates the fixed receptor; the second segment '
    '(len_min\u2013len_max) specifies the desired length range for the diffused scaffold. '
    'Target length ranges were set to 60\u201375% of the native catalytic domain size, '
    'determined by structural analysis of each enzyme\u2019s domain architecture. '
    'For example, lysozyme (129 residues total) used the contig [A35-52/0 72-98], '
    'fixing residues 35\u201352 (18 residues containing Glu35 and Asp52) '
    'and targeting a scaffold of 72\u201398 residues. '
    'The complete contig specifications for all eight enzymes are provided in Supplementary Table S1.'
)

P(
    'For Pseudomonas cepacia lipase, the catalytic triad (Ser87, Asp264, His286) spans 200 residues '
    'in primary sequence, making a single contiguous motif strategy impractical: '
    'fixing residues 85\u2013268 would lock 184 residues (58% of the enzyme), '
    'precluding meaningful size reduction. We therefore employed a multi-motif strategy '
    'in which three short segments of 5\u20137 residues each are independently fixed around each catalytic residue, '
    'separated by flexible linkers of 5\u201315 residues. The contig specification was '
    '[A85-89/0 5-15/0 A262-266/0 5-15/0 A284-288/0 100-150]. '
    'This reduces the total fixed footprint from 184 residues to 15 residues (5%), '
    'leaving 95% of the sequence space available for diffusion-driven compaction.'
)

P(
    'All backbone generation used the base RFdiffusion model checkpoint '
    '(available at https://github.com/RosettaCommons/RFdiffusion) with the following parameters: '
    'diffusion steps T = 50, linear \u03b2 schedule with \u03b2\u2080 = 0.01 and \u03b2T = 0.07, '
    'coordinate scale c_std = 0.25, and IGSO3 rotation schedule. '
    'These are the default parameters distributed with the base model and were used without modification. '
    'We generated 200 backbone candidates per enzyme per strategy, '
    'for a total of 1,800 backbone PDB files (7 \u00d7 200 single-motif '
    '+ 200 single-motif pc_lipase for comparison + 200 multi-motif pc_lipase). '
    'A motif RMSD filter of < 1.0 \u00c5 was applied at generation time; '
    'designs exceeding this threshold were discarded and regenerated. '
    'All RFdiffusion runs were performed on a desktop workstation '
    'equipped with an NVIDIA RTX 4070 Ti SUPER GPU (16 GB VRAM). '
    'Per-design generation times ranged from approximately 0.34 minutes (lysozyme, 85-residue target scaffold) '
    'to approximately 1.2 minutes (glucose oxidase, 350-residue target scaffold). '
    'The complete set of 1,800 backbones required approximately 24 hours of GPU time.'
)

# === 2.3 Active Site Geometry Analysis ===
H('2.3 Active Site Geometry Analysis', 2)

P(
    'To quantitatively assess whether miniaturized backbones preserve the spatial arrangement '
    'of catalytic residues, we performed post-hoc geometry analysis on all 1,800 RFdiffusion designs. '
    'For each design, we extracted the C\u03b1 coordinates of all catalytic residues from the fixed motif chain '
    'and computed the root-mean-square deviation (RMSD) relative to the corresponding residues '
    'in the native crystal structure. Prior to RMSD computation, the designed and native structures '
    'were superimposed to achieve optimal alignment. For motifs containing three or more catalytic residues, '
    'we used the Kabsch algorithm to find the rotation matrix minimizing the sum of squared distances. '
    'For motifs with only two catalytic residues (lysozyme, triosephosphate isomerase, glucose oxidase, '
    'tropinone reductase-II), Kabsch alignment is underdetermined; we instead computed centroid-based '
    'pairwise distance RMSD after translating both coordinate sets to their respective centers of mass.'
)

P(
    'We report five statistics for each enzyme: mean catalytic RMSD, standard deviation, '
    'RMSD range (minimum to maximum), and the fraction of designs below two biologically meaningful thresholds: '
    '< 0.5 \u00c5 (well-preserved geometry, within the range of typical crystallographic coordinate error) '
    'and < 1.0 \u00c5 (adequately preserved, sub-angstrom). '
    'These thresholds provide complementary information: the 0.5 \u00c5 threshold identifies designs '
    'with near-native catalytic geometry suitable for high-confidence applications, '
    'while the 1.0 \u00c5 threshold serves as a minimum quality gate for further evaluation.'
)

P(
    'To test the relationship between backbone geometry quality and sequence design quality, '
    'we computed the Pearson correlation coefficient between catalytic RMSD '
    'and the best ProteinMPNN global score for each backbone (see Section 2.4). '
    'A correlation near zero would indicate that geometry and sequence quality are independent dimensions, '
    'while a strong negative correlation would suggest that better geometry enables better sequence design. '
    'All geometry analyses were performed using custom Python scripts '
    'leveraging BioPython\u2019s PDB parser and NumPy for linear algebra operations. '
    'The analysis code is included in the public repository.'
)

# === 2.4 ProteinMPNN Sequence Design ===
H('2.4 ProteinMPNN Sequence Design', 2)

P(
    'For each enzyme, we selected the 50 top-ranked RFdiffusion backbones '
    '(ranked by catalytic motif RMSD, lowest first) for sequence design. '
    'We used ProteinMPNN version v_48_020, the same version described in the original publication '
    '(Dauparas et al., 2022, Science), accessed from the official repository '
    'at https://github.com/dauparas/ProteinMPNN. All sequence design runs used '
    'C\u03b1-only mode (backbone atom representation without side chains), '
    'a sampling temperature of T = 0.2, and 8 independent sequences generated per backbone. '
    'The temperature of 0.2 was chosen as a moderate value that balances sequence diversity '
    'with structural compatibility: lower temperatures (0.1) produce more conservative designs '
    'with higher native sequence recovery, while higher temperatures (0.3\u20130.5) '
    'explore more diverse sequence space at the risk of reduced foldability. '
    'We did not apply backbone Gaussian noise (the augment_eps parameter was left at its default), '
    'as the RFdiffusion backbones already represent distinct structural solutions '
    'and additional conformational sampling was not required.'
)

P(
    'A critical design decision concerns which protein chains are fixed and which are designed. '
    'In our RFdiffusion output, the fixed catalytic motif occupies chain A '
    '(and chains B and C for the multi-motif pc_lipase designs), '
    'while the diffused scaffold occupies chain B (or chain D). '
    'We specified the design chains using ProteinMPNN\u2019s --pdb_path_chains argument, '
    'which constrains sequence design to the specified chains while preserving '
    'the native residue identities of all other chains. This chain-specific locking '
    'ensures that catalytic residues retain their exact wild-type identities, '
    'while the surrounding scaffold residues are freely optimized for folding compatibility '
    'with the RFdiffusion-generated backbone. Failure to apply this constraint '
    'would result in ProteinMPNN redesigning the catalytic motif as well as the scaffold, '
    'potentially disrupting catalytic function. We verified correct chain assignment '
    'by inspecting the fixed_chains and designed_chains annotations '
    'in each output FASTA header.'
)

P(
    'Each ProteinMPNN run produced 8 designed sequences per backbone, '
    'each annotated with a global sequence design score S_global '
    '(lower scores indicate better sequence\u2013structure compatibility, '
    'computed as the mean per-residue negative log-likelihood under the model) '
    'and per-residue sequence recovery relative to the input backbone. '
    'The total number of designed sequences is 4,800: '
    '7 enzymes \u00d7 50 backbones \u00d7 8 sequences = 2,800 from the single-motif benchmark, '
    'plus 50 backbones \u00d7 8 sequences = 400 from the single-motif pc_lipase comparison, '
    'plus 200 backbones \u00d7 8 sequences = 1,600 from the multi-motif pc_lipase strategy. '
    'All sequences were saved in standardized FASTA format '
    'with score and recovery annotations in the header lines.'
)

# === 2.5 ESM-2 Embedding Analysis ===
H('2.5 ESM-2 Embedding Analysis with Three-Way Comparison', 2)

P(
    'To characterize the sequence space occupied by computationally designed mini-enzymes, '
    'we extracted protein language model embeddings for all 4,800 designed sequences '
    'using ESM-2 (Lin et al., 2023, Science), specifically the esm2_t30_150M_UR50D model '
    '(150 million parameters, 30 transformer layers, 640-dimensional embeddings). '
    'ESM-2 was chosen because it is the most widely used protein language model for '
    'sequence representation tasks, it is openly available, and its 150M parameter variant '
    'runs efficiently on consumer-grade GPUs. Embeddings were extracted from layer 30, '
    'which ESM-2 studies have shown to capture the most informative representations '
    'for downstream structure and function prediction tasks. '
    'For each designed sequence, we computed a single embedding vector '
    'as the mean token representation over all amino acid positions in the scaffold chain, '
    'excluding the fixed catalytic motif residues. This mean-pooling approach '
    'produces a fixed-length 640-dimensional vector that summarizes the global '
    'sequence characteristics of each design.'
)

P(
    'We constructed a three-way comparison framework to contextualize the embedding-space '
    'behavior of designed sequences against two baselines: '
    '(1) Intra-design similarity (Sim_intra): mean pairwise cosine similarity among '
    'all 400 designed sequences for a given enzyme (8 sequences per backbone \u00d7 50 backbones), '
    'measuring the sequence diversity generated by ProteinMPNN for different backbone solutions. '
    '(2) Native-design similarity (Sim_native): cosine similarity between the native wild-type '
    'enzyme sequence (mean-pooled ESM-2 embedding from the full-length PDB structure) '
    'and each of the 400 designed sequences, quantifying how far computational redesign '
    'departs from natural evolutionary sequence space. '
    '(3) Random-sequence baseline (Sim_rand): pairwise cosine similarity among '
    '200 random amino acid sequences whose lengths are drawn from the same distribution '
    'as our designed scaffolds (80\u2013400 residues). Random sequences were generated '
    'by sampling uniformly from the 20 standard amino acids. '
    'This baseline provides an empirically derived null distribution '
    'against which the structure and order of designed sequences can be assessed. '
    'All ESM-2 inference was performed in PyTorch with half-precision floating point '
    'on the same NVIDIA RTX 4070 Ti SUPER GPU used for backbone generation.'
)

# === 2.6 Evaluation Metrics ===
H('2.6 Evaluation Metrics', 2)

P(
    'We define five quantitative metrics for enzyme miniaturization evaluation. '
    'All metrics are reported as mean \u00b1 standard deviation unless otherwise noted.'
)

P(
    'Size Reduction Efficiency (SRE). SRE = (1 \u2212 N\u0304_designed / N_native) \u00d7 100%, '
    'where N\u0304_designed is the mean total residue count measured from the 200 RFdiffusion-generated '
    'PDB files for a given enzyme, and N_native is the residue count of the native enzyme '
    '(taken from the PDB structure, counting standard residues only). '
    'Positive SRE values indicate successful size reduction; zero indicates no change; '
    'negative values indicate the design is larger than the native enzyme. '
    'The total residue count includes both the fixed motif chain(s) and the designed scaffold chain.'
)

P(
    'Catalytic Residue C\u03b1 RMSD. The root-mean-square deviation of catalytic residue '
    'C\u03b1 atom positions after optimal superposition of the designed structure onto the native structure, '
    'as described in Section 2.3. Reported as mean \u00b1 SD across 200 designs per enzyme.'
)

P(
    'ProteinMPNN Global Score (S_global). The global sequence design score reported by ProteinMPNN, '
    'computed as the mean per-residue negative log-likelihood under the inverse folding model. '
    'Lower scores indicate better predicted sequence\u2013structure compatibility. '
    'For each enzyme, we report the best (lowest) score among all 400 designed sequences '
    'and the mean \u00b1 SD across all 400 sequences.'
)

P(
    'ESM-2 Intra-Enzyme Similarity (Sim_intra). The mean pairwise cosine similarity '
    'among all 400 designed sequences (8 sequences per backbone \u00d7 50 backbones) for a given enzyme, '
    'excluding self-comparisons. A value of 1.0 would indicate all designs produce identical embeddings; '
    'lower values indicate greater sequence diversity.'
)

P(
    'ESM-2 Native-Design Similarity (Sim_native). The mean cosine similarity between the native '
    'wild-type enzyme embedding and each of the 400 designed sequence embeddings. '
    'A value near 1.0 would indicate that designed sequences closely resemble the native enzyme; '
    'lower values indicate greater divergence from the natural sequence.'
)

# === 2.7 Statistical Methods ===
H('2.7 Statistical Methods', 2)

P(
    'All reported values are mean \u00b1 standard deviation unless otherwise noted. '
    'For the catalytic RMSD versus ProteinMPNN score independence analysis, '
    'we computed Pearson correlation coefficients per enzyme using the 50 paired observations '
    '(one per backbone). The null hypothesis of zero correlation was not rejected for any enzyme '
    '(all |r| < 0.22, all p > 0.1 by two-tailed t-test with n = 50). '
    'For the ESM-2 three-way comparison, the random baseline distribution '
    'was characterized by computing pairwise cosine similarities among 200 random sequences '
    '(19,900 pairwise comparisons, excluding self-comparisons). '
    'The significance of the observed intra-design and native-design similarity values '
    'was assessed using z-scores relative to this null distribution. '
    'All statistical analyses were performed using NumPy and SciPy.'
)

# === 2.8 Computational Environment ===
H('2.8 Computational Environment', 2)

P(
    'All computational experiments were performed on a desktop workstation '
    'with the following specifications: Intel Core i7-13700K CPU, '
    '64 GB DDR5 system memory, NVIDIA GeForce RTX 4070 Ti SUPER GPU (16 GB VRAM, CUDA 12.6), '
    'running Windows 11 with Windows Subsystem for Linux (WSL2, Ubuntu 22.04). '
    'The protein_design conda environment included Python 3.11, PyTorch 2.12, '
    'BioPython 1.83, NumPy 1.26, and the respective dependencies for RFdiffusion v1.1, '
    'ProteinMPNN v_48_020, and ESM-2 (via the fair-esm package). '
    'ColabFold 1.6.1 was installed but could not utilize the GPU due to TensorFlow '
    'library incompatibilities in the WSL environment; '
    'structural validation was therefore performed via ESM-2 embedding analysis '
    'rather than AlphaFold2 prediction (see Limitations, Section 4.2).'
)

doc.save(OUT_PATH)
print('Saved:', OUT_PATH)
