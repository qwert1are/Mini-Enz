from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT_PATH = r"H:\eazyclaw\saved\MiniEnz_Methodology\reports\title_V2.docx"

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'; sty.font.size = Pt(11)

# ===== TITLE =====
tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = tt.add_run('Multi-Motif Scaffolding Enables Deep Learning-Guided Enzyme Miniaturization:\nA Pilot Benchmark Across Eight Structurally Diverse Enzymes')
rr.bold = True; rr.font.size = Pt(16); rr.font.name = 'Times New Roman'

doc.add_paragraph()

# ===== AUTHOR =====
au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
au.add_run('[Author names withheld for double-blind review]').font.size = Pt(10)
doc.add_paragraph()

# ===== ABSTRACT =====
doc.add_heading('Abstract', 1)
doc.add_paragraph(
    'Can deep generative models reliably shrink enzymes while preserving their catalytic machinery? '
    'We address this question through MiniEnz, a computational framework that evaluates enzyme miniaturization '
    'across eight structurally diverse enzymes spanning six EC classes, seven fold types, '
    'and a 4.5-fold range in protein length (129 to 581 residues). Our three-stage pipeline '
    'combines RFdiffusion backbone generation with a novel multi-motif scaffolding strategy '
    '(1,800 scaffolds), ProteinMPNN sequence design with chain-specific constraints (4,800 sequences), '
    'and ESM-2 embedding analysis within a three-way comparison framework. '
    'Size reduction ranged from 6% for a TIM barrel enzyme to 50% for a bacterial lipase, '
    'with all 1,800 designs preserving sub-angstrom catalytic geometry. '
    'A central finding is that catalytic geometry and sequence design quality are statistically independent '
    'evaluation dimensions, showing near-zero correlation across all enzymes. '
    'We further demonstrate that multi-motif scaffolding\u2014fixing only catalytic residues as short independent '
    'segments rather than a single contiguous block\u2014improves size reduction by 70 percentage points '
    'for enzymes with dispersed active sites. '
    'We classify three systematic failure modes and provide actionable design guidelines. '
    'MiniEnz establishes the first reusable evaluation framework for computational enzyme miniaturization. '
    'We publicly release all 1,800 backbone scaffolds, 4,800 designed sequences, and analysis code.')

kw = doc.add_paragraph()
kw.add_run('Keywords: ').bold = True
kw.add_run('Protein miniaturization \u00b7 RFdiffusion \u00b7 ProteinMPNN \u00b7 '
           'ESM-2 embeddings \u00b7 Motif scaffolding \u00b7 Computational enzyme design').font.size = Pt(10)

# ===== COMPARISON TABLE =====
doc.add_heading('Title, Abstract, and Keywords \u2014 Optimization Notes', 2)

T = lambda hd, data: doc.add_table(rows=1+len(data), cols=len(hd)) or [
    setattr(doc.tables[-1], 'style', 'Light List Accent 1'),
    [setattr(doc.tables[-1].cell(0,j).paragraphs[0].runs[0] if doc.tables[-1].cell(0,j).paragraphs[0].runs else None, 
             '__dummy__', None) for j in range(len(hd))],
    [setattr(doc.tables[-1].cell(i+1,j), 'text', str(val)) for i,row in enumerate(data) for j,val in enumerate(row)]
]

tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Light List Accent 1'
headers = ['Element', 'Previous Version', 'Optimized Version']
data = [
    ['Title',
     'Computational Enzyme Miniaturization at Scale: A Multi-Motif Framework and Systematic Pilot Benchmark',
     'Multi-Motif Scaffolding Enables Deep Learning-Guided Enzyme Miniaturization: A Pilot Benchmark Across Eight Structurally Diverse Enzymes'],
    ['Abstract opening',
     'Deep generative models have transformed protein design, yet enzyme miniaturization remains an unsolved challenge...',
     'Can deep generative models reliably shrink enzymes while preserving their catalytic machinery?'],
    ['Keywords (before)',
     'Enzyme miniaturization \u00b7 Diffusion models \u00b7 ProteinMPNN \u00b7 Protein language models \u00b7 Benchmark \u00b7 De novo protein design',
     'Protein miniaturization \u00b7 RFdiffusion \u00b7 ProteinMPNN \u00b7 ESM-2 embeddings \u00b7 Motif scaffolding \u00b7 Computational enzyme design'],
]
for j, h in enumerate(headers):
    c = tbl.cell(0, j); c.text = h
    for p in c.paragraphs:
        for r in p.runs: r.bold = True; r.font.size = Pt(9)
for i, row in enumerate(data):
    for j, val in enumerate(row):
        c = tbl.cell(i+1, j); c.text = val
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(9)

doc.add_paragraph()

# ===== OPTIMIZATION RATIONALE =====
doc.add_heading('Rationale', 2)

rationale = [
    ('Title', 
     'The new title leads with the methodological contribution (multi-motif scaffolding) rather than the '
     'tool name or generic descriptors. It specifies the enzyme count, signaling honesty about pilot scale '
     'while the word "Enables" conveys proven capability rather than aspiration.'),
    ('Abstract',
     'The opening question engages the reader immediately. The single-paragraph flow moves from question '
     'to method to key results to implications without section labels (Purpose/Methods/Results). '
     'The 70-percentage-point improvement is stated as a concrete number, not a vague claim. '
     'The em-dash pair around the multi-motif explanation is the only one in the entire abstract\u2014'
     'used deliberately to define the core methodological concept.'),
    ('Keywords',
     'Replaced "Diffusion models" with "RFdiffusion" for specificity. '
     'Replaced "Protein language models" with "ESM-2 embeddings" to match the actual tool used. '
     'Replaced "De novo protein design" with "Computational enzyme design" since the work is about '
     'redesigning existing enzymes rather than de novo generation. '
     'Added "Motif scaffolding" as a methodological keyword.'),
]

for title, text in rationale:
    p = doc.add_paragraph()
    p.add_run(title + '. ').bold = True
    p.add_run(text)

doc.save(OUT_PATH)
print('Saved:', OUT_PATH)
